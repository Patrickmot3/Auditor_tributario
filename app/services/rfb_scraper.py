"""
Serviço de consulta e atualização das tabelas oficiais SPED/RFB.
Detecta automaticamente o formato do arquivo (DOCX, XLSX ou HTML).
"""
import io
import logging
import re
import time
import zipfile
from datetime import datetime, timezone, date

import requests
from app.extensions import db
from app.models.ncm import NcmTributario, GrupoTributario
from app.models.base_tributaria import LogAtualizacao

logger = logging.getLogger(__name__)

TABELAS_SPED = {
    # 4.3.10 — Monofásicos: Alíquotas Diferenciadas / Bebidas Frias (CST 02/04)
    '4.3.10': {
        'url_download': 'http://sped.rfb.gov.br/arquivo/download/1638',
        'url_show': 'http://sped.rfb.gov.br/arquivo/show/1638',
        'nome': 'Monofásicos — Alíquotas Diferenciadas (CST 02/04)',
    },
    # 4.3.11 — Monofásicos: Alíquotas por Unidade de Medida / Bebidas Frias (CST 03/04)
    '4.3.11': {
        'url_download': 'http://sped.rfb.gov.br/arquivo/download/5786',
        'url_show': 'http://sped.rfb.gov.br/arquivo/show/5786',
        'nome': 'Monofásicos por Unidade de Medida (CST 03/04)',
    },
    # 4.3.12 — Substituição Tributária (CST 05)
    '4.3.12': {
        'url_download': 'http://sped.rfb.gov.br/arquivo/download/1642',
        'url_show': 'http://sped.rfb.gov.br/arquivo/show/1642',
        'nome': 'Substituição Tributária (CST 05)',
    },
    # 4.3.13 — Fármacos e Perfumaria / Alíquota Zero (CST 06)
    '4.3.13': {
        'url_download': 'http://sped.rfb.gov.br/arquivo/download/1643',
        'url_show': 'http://sped.rfb.gov.br/arquivo/show/1643',
        'nome': 'Alíquota Zero — Fármacos e Perfumaria (CST 06)',
    },
    # 4.3.14 — Isenção PIS/COFINS (CST 07)
    '4.3.14': {
        'url_download': 'http://sped.rfb.gov.br/arquivo/download/1646',
        'url_show': 'http://sped.rfb.gov.br/arquivo/show/1646',
        'nome': 'Isenção PIS/COFINS (CST 07)',
    },
    # 4.3.15 — Sem Incidência (CST 08)
    '4.3.15': {
        'url_download': 'http://sped.rfb.gov.br/arquivo/download/1651',
        'url_show': 'http://sped.rfb.gov.br/arquivo/show/1651',
        'nome': 'Sem Incidência (CST 08)',
    },
    # 4.3.16 — Suspensão PIS/COFINS (CST 09)
    '4.3.16': {
        'url_download': 'http://sped.rfb.gov.br/arquivo/download/1655',
        'url_show': 'http://sped.rfb.gov.br/arquivo/show/1655',
        'nome': 'Suspensão PIS/COFINS (CST 09)',
    },
}

HEADERS = {
    'User-Agent': 'Mozilla/5.0 (compatible; TribSync/1.0)',
}
MAX_TENTATIVAS = 3
INTERVALO_RETRY = 5


# ─── HTTP ─────────────────────────────────────────────────────────────────────

def _get_com_retry(url, tentativas=MAX_TENTATIVAS):
    for i in range(tentativas):
        try:
            resp = requests.get(url, headers=HEADERS, timeout=30)
            resp.raise_for_status()
            return resp
        except requests.RequestException as e:
            logger.warning(f'Tentativa {i+1}/{tentativas} falhou: {url} — {e}')
            if i < tentativas - 1:
                time.sleep(INTERVALO_RETRY)
    raise Exception(f'Falha ao acessar {url} após {tentativas} tentativas')


# ─── Detecção de formato ──────────────────────────────────────────────────────

_OLE2_MAGIC = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'  # Assinatura XLS/DOC antigo (BIFF8)


def _detectar_formato(conteudo: bytes, content_type: str) -> str:
    """
    Detecta o formato do arquivo retornado pela RFB.
    Ordem de confiança: magic bytes > [Content_Types].xml > Content-Type > texto.
    Retorna: 'xls', 'xlsx', 'docx', 'html' ou 'desconhecido'.
    """
    ct = content_type.lower()

    # OLE2 Compound Document (magic D0 CF 11 E0) → XLS ou DOC antigo
    if conteudo[:8] == _OLE2_MAGIC:
        # Distinguir XLS de DOC pelo Content-Type quando possível
        if 'word' in ct or 'msword' in ct:
            return 'doc_antigo'  # tratado como xls no fallback
        return 'xls'

    # Arquivo ZIP (magic bytes PK) → DOCX ou XLSX
    if conteudo[:2] == b'PK':
        try:
            with zipfile.ZipFile(io.BytesIO(conteudo)) as z:
                if '[Content_Types].xml' in z.namelist():
                    ct_xml = z.read('[Content_Types].xml').decode('utf-8', errors='ignore')
                    if 'wordprocessingml' in ct_xml:
                        return 'docx'
                    if 'spreadsheetml' in ct_xml:
                        return 'xlsx'
                nomes = z.namelist()
                if any('xl/worksheets' in n or 'xl/workbook' in n for n in nomes):
                    return 'xlsx'
                if any('word/document' in n for n in nomes):
                    return 'docx'
        except Exception as e:
            logger.warning(f'Erro ao inspecionar ZIP: {e}')

    # Content-Type do servidor
    if 'spreadsheet' in ct or 'excel' in ct or 'xlsx' in ct:
        return 'xlsx'
    if 'ms-excel' in ct or ct.endswith('/xls'):
        return 'xls'
    if 'word' in ct or 'docx' in ct or 'msword' in ct:
        return 'docx'
    if 'html' in ct or 'text/plain' in ct:
        return 'html'

    # Último recurso: texto
    try:
        texto = conteudo[:200].decode('utf-8', errors='ignore').lower()
        if '<html' in texto or '<!doctype' in texto:
            return 'html'
    except Exception:
        pass

    return 'desconhecido'


# ─── Parsers ──────────────────────────────────────────────────────────────────

def _extrair_ncms_docx(conteudo: bytes) -> list[tuple[str, str]]:
    """Extrai pares (ncm, descricao) de arquivo DOCX."""
    from docx import Document
    doc = Document(io.BytesIO(conteudo))
    resultado = []
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells]
            if len(celulas) < 2:
                continue
            ncm_raw = celulas[0].replace('.', '').replace('-', '').strip()
            if re.match(r'^\d{4,10}$', ncm_raw):
                resultado.append((ncm_raw, celulas[1]))
    # Também varrer parágrafos para NCMs em texto corrido
    for para in doc.paragraphs:
        match = re.match(r'^(\d[\d.]{3,})\s+(.*)', para.text.strip())
        if match:
            ncm_raw = match.group(1).replace('.', '').replace('-', '').strip()
            if re.match(r'^\d{4,10}$', ncm_raw):
                resultado.append((ncm_raw, match.group(2).strip()))
    return resultado


def _extrair_ncms_xlsx(conteudo: bytes) -> list[tuple[str, str]]:
    """Extrai pares (ncm, descricao) de arquivo XLSX."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(conteudo), read_only=True, data_only=True)
    resultado = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            if not row or row[0] is None:
                continue
            ncm_raw = str(row[0]).replace('.', '').replace('-', '').strip()
            if re.match(r'^\d{4,10}$', ncm_raw):
                descricao = str(row[1]).strip() if len(row) > 1 and row[1] else ''
                resultado.append((ncm_raw, descricao))
    return resultado


def _extrair_ncms_xls(conteudo: bytes) -> list[tuple[str, str]]:
    """Extrai pares (ncm, descricao) de arquivo XLS antigo (BIFF8/OLE2) via xlrd."""
    import xlrd
    wb = xlrd.open_workbook(file_contents=conteudo)
    resultado = []
    for ws in wb.sheets():
        for rx in range(ws.nrows):
            row = ws.row(rx)
            if not row:
                continue
            cel0 = row[0]
            # xlrd retorna XL_CELL_FLOAT para números
            if cel0.ctype == xlrd.XL_CELL_FLOAT:
                ncm_raw = str(int(cel0.value))
            elif cel0.ctype == xlrd.XL_CELL_TEXT:
                ncm_raw = cel0.value.replace('.', '').replace('-', '').strip()
            else:
                continue
            if re.match(r'^\d{4,10}$', ncm_raw):
                descricao = ''
                if len(row) > 1 and row[1].ctype == xlrd.XL_CELL_TEXT:
                    descricao = row[1].value.strip()
                resultado.append((ncm_raw, descricao))
    return resultado


def _extrair_ncms_ole2_doc(conteudo: bytes) -> list[tuple[str, str]]:
    """
    Extrai NCMs de arquivo OLE2 Word DOC (binário .doc antigo) via olefile.
    Lê todos os streams internos e aplica regex sobre texto decodificado.
    Word 97-2003 armazena texto em UTF-16 LE nos streams internos.
    """
    import olefile

    ole = olefile.OleFileIO(io.BytesIO(conteudo))
    resultado = []
    vistos = set()

    def _buscar(texto: str):
        texto = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', ' ', texto)
        texto = re.sub(r' {4,}', '   ', texto)
        for m in re.finditer(
            r'(\d[\d.]{3,})\s{1,10}([A-ZÀ-Üa-zà-ü][^\n\r\x00]{3,150})',
            texto,
        ):
            ncm_raw = m.group(1).replace('.', '').replace('-', '').strip()
            if re.match(r'^\d{4,10}$', ncm_raw) and ncm_raw not in vistos:
                desc = m.group(2).strip()[:300]
                resultado.append((ncm_raw, desc))
                vistos.add(ncm_raw)

    for entry in ole.listdir():
        try:
            data = ole.openstream(entry).read()
            if len(data) < 20:
                continue
            # UTF-16 LE — padrão em Word 97+ (Unicode)
            _buscar(data.decode('utf-16-le', errors='ignore'))
            # CP1252 — Word antigo não-Unicode
            _buscar(data.decode('cp1252', errors='ignore'))
        except Exception:
            continue

    # Fallback: varrer bytes brutos do arquivo inteiro
    if not resultado:
        _buscar(conteudo.decode('utf-16-le', errors='ignore'))
        _buscar(conteudo.decode('cp1252', errors='ignore'))

    return resultado


def _extrair_ncms_html(conteudo: bytes) -> list[tuple[str, str]]:
    """Extrai pares (ncm, descricao) de HTML com tabelas."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(conteudo, 'lxml')
    resultado = []
    for tr in soup.find_all('tr'):
        tds = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
        if len(tds) < 2:
            continue
        ncm_raw = tds[0].replace('.', '').replace('-', '').strip()
        if re.match(r'^\d{4,10}$', ncm_raw):
            resultado.append((ncm_raw, tds[1]))
    return resultado


def _extrair_ncms_zip_xml(conteudo: bytes) -> list[tuple[str, str]]:
    """
    Fallback: lê XMLs diretamente do ZIP sem depender de python-docx/openpyxl.
    Funciona para OOXML não-padrão (ex.: tabelas SPED com Content-Types incomum).
    Estratégia:
      1. Prioriza xl/sharedStrings.xml (XLSX) e word/document.xml (DOCX).
      2. Varre todos os .xml do ZIP caso esses não existam.
      3. Aplica regex para encontrar padrões NCM seguidos de descrição.
    """
    from bs4 import BeautifulSoup

    if conteudo[:2] != b'PK':
        logger.debug('_extrair_ncms_zip_xml: conteúdo não é ZIP, abortando silenciosamente')
        return []

    resultado = []
    vistos = set()

    def _extrair_de_xml(xml_bytes: bytes):
        """Extrai pares NCM+descrição de um bloco XML."""
        soup = BeautifulSoup(xml_bytes, 'lxml-xml')
        textos = [t.strip() for t in soup.stripped_strings if t.strip()]

        i = 0
        while i < len(textos):
            ncm_raw = textos[i].replace('.', '').replace('-', '').strip()
            if re.match(r'^\d{4,10}$', ncm_raw) and ncm_raw not in vistos:
                descricao = textos[i + 1].strip() if i + 1 < len(textos) else ''
                # Ignora se a "descrição" também parece ser um NCM
                if re.match(r'^\d{4,10}$', descricao.replace('.', '').replace('-', '')):
                    descricao = ''
                resultado.append((ncm_raw, descricao))
                vistos.add(ncm_raw)
                i += 2
                continue
            i += 1

    # Também tenta varredura linha-a-linha com regex (para texto corrido)
    def _extrair_de_texto(xml_bytes: bytes):
        texto = xml_bytes.decode('utf-8', errors='ignore')
        # Remove tags XML para obter texto limpo
        texto_limpo = re.sub(r'<[^>]+>', ' ', texto)
        for match in re.finditer(
            r'\b(\d[\d.]{3,})\s{1,10}([A-ZÀ-Ü][^\n]{3,120})',
            texto_limpo,
        ):
            ncm_raw = match.group(1).replace('.', '').replace('-', '').strip()
            if re.match(r'^\d{4,10}$', ncm_raw) and ncm_raw not in vistos:
                resultado.append((ncm_raw, match.group(2).strip()))
                vistos.add(ncm_raw)

    PRIORIDADE = [
        'xl/sharedStrings.xml',
        'word/document.xml',
        'xl/worksheets/sheet1.xml',
    ]

    with zipfile.ZipFile(io.BytesIO(conteudo)) as z:
        nomes = z.namelist()
        # Processar arquivos prioritários primeiro
        for nome in PRIORIDADE:
            if nome in nomes:
                try:
                    xml_bytes = z.read(nome)
                    _extrair_de_xml(xml_bytes)
                    _extrair_de_texto(xml_bytes)
                    logger.debug(f'_extrair_ncms_zip_xml: processado {nome} → {len(resultado)} NCMs até agora')
                except Exception as e:
                    logger.debug(f'_extrair_ncms_zip_xml: erro em {nome} — {e}')

        # Se não encontrou nada, varrer todos os XMLs
        if not resultado:
            for nome in nomes:
                if not nome.endswith('.xml') or nome in PRIORIDADE:
                    continue
                try:
                    xml_bytes = z.read(nome)
                    _extrair_de_xml(xml_bytes)
                    _extrair_de_texto(xml_bytes)
                except Exception as e:
                    logger.debug(f'_extrair_ncms_zip_xml: erro em {nome} — {e}')

    return resultado


# Capítulos TIPI válidos por código de tabela SPED (regime monofásico)
_CAPITULOS_VALIDOS = {
    '4.3.10': {'40', '68', '70', '73', '83', '84', '85', '87', '90', '91', '94'},  # Autopeças
    '4.3.11': {'21', '22', '39', '70', '73', '76'},  # Bebidas Frias (por unidade de medida)
    '4.3.13': {'30', '33', '34', '35', '38', '40', '48', '84', '85', '87', '89', '90'},  # Fármacos/Perfumaria
    # 4.3.12, 4.3.14, 4.3.15, 4.3.16: abrangem múltiplos capítulos, sem restrição por capítulo
}

# Faixa de anos que não são NCMs válidos (ex-anos extraídos de notas de rodapé)
_ANO_MIN, _ANO_MAX = 1900, 2050


def _ncm_valido(ncm_raw: str, tabela_id: str) -> bool:
    """Retorna True se o NCM extraído parece um código real da TIPI para a tabela."""
    n = len(ncm_raw)
    if n < 4 or n > 8:
        return False

    # Rejeita 4 dígitos que se encaixam em faixa de anos (ex: 2003, 2015, 1998)
    if n == 4:
        try:
            valor = int(ncm_raw)
            if _ANO_MIN <= valor <= _ANO_MAX:
                logger.debug(f'NCM rejeitado (parece ano): {ncm_raw}')
                return False
        except ValueError:
            return False

    # Rejeita 8 dígitos no formato DDMMAAAA (datas como 01042014)
    if n == 8:
        try:
            dd, mm = int(ncm_raw[:2]), int(ncm_raw[2:4])
            ano = int(ncm_raw[4:])
            if 1 <= dd <= 31 and 1 <= mm <= 12 and _ANO_MIN <= ano <= _ANO_MAX:
                logger.debug(f'NCM rejeitado (parece data DD/MM/AAAA): {ncm_raw}')
                return False
        except ValueError:
            pass

    # Valida capítulo TIPI vs tabela esperada
    capitulos_ok = _CAPITULOS_VALIDOS.get(tabela_id)
    if capitulos_ok:
        capitulo = ncm_raw[:2]
        if capitulo not in capitulos_ok:
            logger.debug(f'NCM rejeitado (capítulo {capitulo} inválido para tabela {tabela_id}): {ncm_raw}')
            return False

    return True


# ─── Persistência ─────────────────────────────────────────────────────────────

def _salvar_ncms(pares: list[tuple[str, str]], tabela_id: str, grupo) -> tuple[int, int]:
    from app.models.base_tributaria import AliquotaGrupo

    # Busca alíquota vigente no banco; usa fallback hardcoded apenas se não houver
    aliquota = AliquotaGrupo.vigente_para(grupo.id)
    pis_fab = float(aliquota.pis_fabricante) if aliquota else 1.5
    cofins_fab = float(aliquota.cofins_fabricante) if aliquota else 7.0
    pis_var = float(aliquota.pis_varejista) if aliquota else 0.0
    cofins_var = float(aliquota.cofins_varejista) if aliquota else 0.0

    if not aliquota:
        logger.warning(
            f'Tabela {tabela_id}: nenhuma AliquotaGrupo vigente para grupo {grupo.id}. '
            'Usando valores padrão hardcoded.'
        )

    inseridos = atualizados = rejeitados = 0
    for ncm_raw, descricao in pares:
        if not _ncm_valido(ncm_raw, tabela_id):
            rejeitados += 1
            continue

        tipo_ref = 'ncm_exato' if len(ncm_raw) == 8 else f'posicao_{len(ncm_raw)}'
        existente = NcmTributario.query.filter_by(
            ncm=ncm_raw, grupo_tributario_id=grupo.id,
        ).first()
        if existente:
            existente.descricao = descricao
            existente.updated_at = datetime.now(timezone.utc)
            atualizados += 1
        else:
            db.session.add(NcmTributario(
                ncm=ncm_raw,
                descricao=descricao,
                grupo_tributario_id=grupo.id,
                monofasico=True,
                tipo_referencia=tipo_ref,
                lei=grupo.lei_base,
                cst_entrada='70',
                cst_saida='04',
                cfop_entrada_simples='1102',
                cfop_saida_simples='5102',
                pis_aliquota_fabricante=pis_fab,
                cofins_aliquota_fabricante=cofins_fab,
                pis_aliquota_varejista=pis_var,
                cofins_aliquota_varejista=cofins_var,
                vigencia_inicio=date.today(),
                fonte_url=TABELAS_SPED[tabela_id]['url_download'],
                ativo=True,
            ))
            inseridos += 1

    if rejeitados:
        logger.info(f'Tabela {tabela_id}: {rejeitados} NCMs rejeitados por validação de capítulo/formato')
    db.session.commit()
    return inseridos, atualizados


# ─── API pública ──────────────────────────────────────────────────────────────

def _checar_versao(tabela_id):
    info = TABELAS_SPED.get(tabela_id)
    if not info:
        return None
    try:
        resp = _get_com_retry(info['url_show'])
        match = re.search(r'Vers[aã]o\s*:?\s*([\d.]+)', resp.text, re.IGNORECASE)
        if match:
            return match.group(1)
        match = re.search(r'(\d{2}/\d{2}/\d{4})', resp.text)
        if match:
            return match.group(1)
    except Exception as e:
        logger.error(f'Erro ao checar versão {tabela_id}: {e}')
    return None


def verificar_atualizacao(tabela_id, executado_por='scheduler_auto'):
    versao_rfb = _checar_versao(tabela_id)
    if not versao_rfb:
        return False

    ultimo = LogAtualizacao.query.filter(
        LogAtualizacao.tabela_sped == tabela_id,
        LogAtualizacao.status.in_(['sucesso', 'seed_inicial']),
    ).order_by(LogAtualizacao.data_importacao.desc()).first()

    if ultimo and ultimo.versao == versao_rfb:
        db.session.add(LogAtualizacao(
            tabela_sped=tabela_id, versao=versao_rfb,
            data_importacao=datetime.now(timezone.utc),
            status='sem_alteracoes',
            mensagem=f'Tabela {tabela_id} já está na versão {versao_rfb}',
            executado_por=executado_por,
        ))
        db.session.commit()
        return False
    return True


def atualizar_tabela(tabela_id, executado_por='scheduler_auto'):
    info = TABELAS_SPED.get(tabela_id)
    if not info:
        return {
            'status': 'erro',
            'mensagem': f'Tabela {tabela_id} não possui URL de atualização automática. Mantenha via seed.',
            'inseridos': 0,
            'atualizados': 0,
        }

    versao_rfb = _checar_versao(tabela_id)
    inseridos = atualizados = 0
    status = 'erro'
    mensagem = ''

    try:
        grupo = GrupoTributario.query.filter_by(tabela_sped=tabela_id).first()
        if not grupo:
            raise Exception(f'Grupo tributário para tabela {tabela_id} não encontrado')

        resp = _get_com_retry(info['url_download'])
        conteudo = resp.content
        content_type = resp.headers.get('Content-Type', '')
        fmt = _detectar_formato(conteudo, content_type)

        logger.info(f'Tabela {tabela_id}: formato detectado = {fmt} | Content-Type = {content_type}')

        # Ordem de tentativa: formato detectado primeiro, depois os outros.
        # _extrair_ncms_xls cobre arquivos BIFF8/OLE2 (xls antigo).
        # _extrair_ncms_zip_xml é sempre o último fallback (lê ZIP diretamente).
        _todos = [_extrair_ncms_ole2_doc, _extrair_ncms_xls, _extrair_ncms_xlsx,
                  _extrair_ncms_docx, _extrair_ncms_html, _extrair_ncms_zip_xml]
        ordem = {
            'xls':       [_extrair_ncms_xls, _extrair_ncms_ole2_doc, _extrair_ncms_xlsx, _extrair_ncms_html, _extrair_ncms_zip_xml],
            'doc_antigo':[_extrair_ncms_ole2_doc, _extrair_ncms_xls, _extrair_ncms_xlsx, _extrair_ncms_html, _extrair_ncms_zip_xml],
            'xlsx':      [_extrair_ncms_xlsx, _extrair_ncms_xls, _extrair_ncms_docx, _extrair_ncms_ole2_doc, _extrair_ncms_html, _extrair_ncms_zip_xml],
            'docx':      [_extrair_ncms_docx, _extrair_ncms_xlsx, _extrair_ncms_xls, _extrair_ncms_ole2_doc, _extrair_ncms_html, _extrair_ncms_zip_xml],
            'html':      [_extrair_ncms_html, _extrair_ncms_xlsx, _extrair_ncms_xls, _extrair_ncms_docx, _extrair_ncms_ole2_doc, _extrair_ncms_zip_xml],
        }.get(fmt, _todos)

        pares = []
        ultimo_erro = None
        for fn in ordem:
            try:
                pares = fn(conteudo)
                if pares:
                    logger.info(f'Tabela {tabela_id}: extraídos {len(pares)} NCMs via {fn.__name__}')
                    break
            except Exception as e:
                ultimo_erro = e
                logger.warning(f'Tabela {tabela_id}: {fn.__name__} falhou — {e}')

        if not pares:
            # Log dos primeiros bytes para diagnóstico de formato desconhecido
            preview = conteudo[:400].decode('utf-8', errors='replace').replace('\n', ' ').strip()
            logger.error(f'Tabela {tabela_id}: preview do conteúdo recebido → {preview!r}')
            raise Exception(f'Nenhum NCM extraído da tabela {tabela_id}. Último erro: {ultimo_erro}')

        inseridos, atualizados = _salvar_ncms(pares, tabela_id, grupo)
        status = 'sucesso'
        mensagem = f'Tabela {tabela_id} ({fmt}): {inseridos} inseridos, {atualizados} atualizados'
        logger.info(mensagem)

    except Exception as e:
        status = 'erro'
        mensagem = str(e)
        logger.error(f'Erro ao atualizar tabela {tabela_id}: {e}')
        db.session.rollback()  # sessão pode estar suja; limpar antes de gravar o log

    try:
        db.session.add(LogAtualizacao(
            tabela_sped=tabela_id,
            versao=versao_rfb,
            data_atualizacao_rfb=date.today(),
            data_importacao=datetime.now(timezone.utc),
            status=status,
            registros_inseridos=inseridos,
            registros_atualizados=atualizados,
            mensagem=mensagem,
            executado_por=executado_por,
        ))
        db.session.commit()
    except Exception as log_err:
        logger.error(f'Falha ao gravar log de atualização para {tabela_id}: {log_err}')
        db.session.rollback()

    return {'status': status, 'inseridos': inseridos,
            'atualizados': atualizados, 'mensagem': mensagem}
